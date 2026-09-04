"""Pure-function tests for app/services/export_rendering.py — no DB, no
Celery, no docling dependency (worker_jobs.py itself can't be imported
on this host venv at all, since it transitively imports docling; this
module deliberately has no such dependency so it's fully testable here).

Also exercises the real committed templates (app/templates/exports/*.docx)
through docxtpl end to end, not just the context-building logic in
isolation — confirms the templates themselves stay in sync with the
context shape these functions produce.
"""

from datetime import datetime, timezone
from types import SimpleNamespace

from docx import Document

from app.services import export_rendering, export_templates


def _section(section_type, content_text, order_index, source_item_id=None):
    return SimpleNamespace(
        section_type=section_type, content_text=content_text,
        order_index=order_index, source_item_id=source_item_id,
    )


def _experience_item(id_, *, title, company, start_date=None, end_date=None, current=False):
    return SimpleNamespace(id=id_, title=title, company=company,
                            start_date=start_date, end_date=end_date, current=current)


def _project_item(id_, *, name):
    return SimpleNamespace(id=id_, name=name)


class TestResolveCandidateNameAndContactLine:
    def test_name_present_in_basics(self):
        assert export_rendering.resolve_candidate_name({"name": "Jane Doe"}) == "Jane Doe"

    def test_name_absent_never_fabricated(self):
        """The CV parser never actually populates 'name' today — this
        must degrade to None, never guess a name from anywhere else."""
        assert export_rendering.resolve_candidate_name({"name": None}) is None
        assert export_rendering.resolve_candidate_name({}) is None
        assert export_rendering.resolve_candidate_name(None) is None

    def test_contact_line_from_basics_email(self):
        assert export_rendering.resolve_contact_line({"email": "jane@example.com"}) == "jane@example.com"

    def test_contact_line_falls_back_to_account_email(self):
        """basics.email is always None in this codebase's real CV parser
        today — an account-owned export should still show the user's
        login email as the one real contact detail available."""
        line = export_rendering.resolve_contact_line({"email": None}, fallback_email="jane@example.com")
        assert line == "jane@example.com"

    def test_contact_line_none_when_nothing_available(self):
        """A trial-session export with no basics.email and no account
        email at all — must omit the line entirely, never fabricate."""
        assert export_rendering.resolve_contact_line({}, fallback_email=None) is None

    def test_contact_line_joins_multiple_present_fields(self):
        line = export_rendering.resolve_contact_line(
            {"email": "jane@example.com", "phone": "555-0199", "location": "London"}
        )
        assert line == "jane@example.com | 555-0199 | London"


class TestBuildCvDocxContext:
    def test_groups_multiple_experience_sections_under_one_heading(self):
        exp1 = _experience_item("e1", title="Senior Designer", company="Voy Health")
        exp2 = _experience_item("e2", title="Designer", company="Acme")
        sections = [
            _section("summary", "Product designer with 5 years experience.", 0),
            _section("experience", "Led the portal redesign.", 1, source_item_id="e1"),
            _section("experience", "Shipped the onboarding flow.", 2, source_item_id="e2"),
            _section("skills", "Figma, Prototyping", 3),
        ]
        ctx = export_rendering.build_cv_docx_context(
            candidate_name=None, contact_line=None, sections=sections,
            experience_by_id={"e1": exp1, "e2": exp2}, project_by_id={},
        )
        exp_blocks = [b for b in ctx["blocks"] if b["kind"] == "experience"]
        assert len(exp_blocks) == 1, "Both experience sections must collapse into one heading block"
        assert len(exp_blocks[0]["entries"]) == 2
        assert exp_blocks[0]["entries"][0]["header"] == "Senior Designer — Voy Health"
        assert exp_blocks[0]["entries"][1]["header"] == "Designer — Acme"

    def test_experience_header_includes_date_range(self):
        exp = _experience_item(
            "e1", title="Senior Designer", company="Voy Health",
            start_date=datetime(2021, 4, 1, tzinfo=timezone.utc), current=True,
        )
        sections = [_section("experience", "Led the redesign.", 0, source_item_id="e1")]
        ctx = export_rendering.build_cv_docx_context(
            candidate_name=None, contact_line=None, sections=sections,
            experience_by_id={"e1": exp}, project_by_id={},
        )
        header = ctx["blocks"][0]["entries"][0]["header"]
        assert "Apr 2021" in header
        assert "Present" in header

    def test_unresolvable_source_item_id_omits_header_not_bullets(self):
        """source_item_id pointing at a row that isn't in the lookup dict
        (deleted since generation, or predates the column) must never
        block rendering — the bullets still show, just without a header."""
        sections = [_section("experience", "Led the redesign.", 0, source_item_id="does-not-exist")]
        ctx = export_rendering.build_cv_docx_context(
            candidate_name=None, contact_line=None, sections=sections,
            experience_by_id={}, project_by_id={},
        )
        entry = ctx["blocks"][0]["entries"][0]
        assert entry["header"] is None
        assert entry["bullets"] == ["Led the redesign."]

    def test_multiline_content_text_splits_into_multiple_bullets(self):
        sections = [_section("experience", "Bullet one.\nBullet two.\n", 0, source_item_id="e1")]
        ctx = export_rendering.build_cv_docx_context(
            candidate_name=None, contact_line=None, sections=sections,
            experience_by_id={"e1": _experience_item("e1", title="X", company="Y")},
            project_by_id={},
        )
        assert ctx["blocks"][0]["entries"][0]["bullets"] == ["Bullet one.", "Bullet two."]

    def test_leading_bullet_glyphs_stripped(self):
        sections = [_section("experience", "• Led the redesign.", 0, source_item_id="e1")]
        ctx = export_rendering.build_cv_docx_context(
            candidate_name=None, contact_line=None, sections=sections,
            experience_by_id={"e1": _experience_item("e1", title="X", company="Y")},
            project_by_id={},
        )
        assert ctx["blocks"][0]["entries"][0]["bullets"] == ["Led the redesign."]

    def test_education_section_becomes_line_list(self):
        sections = [_section("education", "BSc Product Design — Manchester (2016)\nAWS Cert (2020)", 0)]
        ctx = export_rendering.build_cv_docx_context(
            candidate_name=None, contact_line=None, sections=sections,
            experience_by_id={}, project_by_id={},
        )
        assert ctx["blocks"][0]["kind"] == "education"
        assert ctx["blocks"][0]["lines"] == ["BSc Product Design — Manchester (2016)", "AWS Cert (2020)"]
        assert ctx["blocks"][0]["paragraph"] is None

    def test_projects_grouped_and_headered_by_name(self):
        proj = _project_item("p1", name="Telehealth Booking Redesign")
        sections = [_section("projects", "Reduced drop-off by 30%.", 0, source_item_id="p1")]
        ctx = export_rendering.build_cv_docx_context(
            candidate_name=None, contact_line=None, sections=sections,
            experience_by_id={}, project_by_id={"p1": proj},
        )
        proj_block = ctx["blocks"][0]
        assert proj_block["kind"] == "projects"
        assert proj_block["entries"][0]["header"] == "Telehealth Booking Redesign"

    def test_sections_ordered_by_order_index_regardless_of_input_order(self):
        sections = [
            _section("skills", "Figma", 2),
            _section("summary", "A summary.", 0),
            _section("education", "A degree.", 1),
        ]
        ctx = export_rendering.build_cv_docx_context(
            candidate_name=None, contact_line=None, sections=sections,
            experience_by_id={}, project_by_id={},
        )
        assert [b["kind"] for b in ctx["blocks"]] == ["summary", "education", "skills"]


class TestBuildCoverLetterDocxContext:
    def test_splits_body_text_on_blank_lines(self):
        ctx = export_rendering.build_cover_letter_docx_context(
            candidate_name="Jane Doe", contact_line="jane@example.com",
            sent_date="August 12, 2026", employer_name="HealthTech Co",
            body_text="Dear Hiring Manager,\n\nI am writing to apply.\n\nSincerely, Jane",
        )
        assert ctx["body_paragraphs"] == [
            "Dear Hiring Manager,", "I am writing to apply.", "Sincerely, Jane",
        ]

    def test_single_paragraph_body_with_no_blank_lines_still_shown(self):
        ctx = export_rendering.build_cover_letter_docx_context(
            candidate_name=None, contact_line=None, sent_date="August 12, 2026",
            employer_name=None, body_text="One unbroken paragraph.",
        )
        assert ctx["body_paragraphs"] == ["One unbroken paragraph."]


class TestRenderDocxAgainstRealTemplates:
    """Integration-style, but still fully host-testable (no docling) —
    proves the real committed template files actually stay in sync with
    the context shape these functions produce, not just that the
    functions build a plausible-looking dict."""

    def test_cv_standard_template_renders_and_contains_expected_text(self):
        ctx = export_rendering.build_cv_docx_context(
            candidate_name="Jane Doe", contact_line="jane@example.com",
            sections=[
                _section("summary", "Product designer with 5 years experience.", 0),
                _section("experience", "Led the portal redesign.", 1, source_item_id="e1"),
                _section("skills", "Figma, Prototyping", 2),
            ],
            experience_by_id={"e1": _experience_item("e1", title="Senior Designer", company="Voy Health")},
            project_by_id={},
        )
        path = export_templates.cv_template_path(export_templates.DEFAULT_CV_TEMPLATE_ID)
        file_bytes = export_rendering.render_docx(path, ctx)
        assert len(file_bytes) > 0

        from io import BytesIO
        doc = Document(BytesIO(file_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Jane Doe" in full_text
        assert "Product designer with 5 years experience." in full_text
        assert "Senior Designer — Voy Health" in full_text
        assert "Led the portal redesign." in full_text
        assert "Figma, Prototyping" in full_text

    def test_cv_compact_template_renders(self):
        ctx = export_rendering.build_cv_docx_context(
            candidate_name=None, contact_line=None,
            sections=[_section("summary", "A summary.", 0)],
            experience_by_id={}, project_by_id={},
        )
        path = export_templates.cv_template_path("compact")
        file_bytes = export_rendering.render_docx(path, ctx)
        assert len(file_bytes) > 0

    def test_cv_template_renders_with_no_name_and_no_contact(self):
        """The honest common case given this parser never extracts a
        name/email/phone/location — must render cleanly, not crash."""
        ctx = export_rendering.build_cv_docx_context(
            candidate_name=None, contact_line=None,
            sections=[_section("summary", "A summary with no header info.", 0)],
            experience_by_id={}, project_by_id={},
        )
        path = export_templates.cv_template_path(export_templates.DEFAULT_CV_TEMPLATE_ID)
        file_bytes = export_rendering.render_docx(path, ctx)
        assert len(file_bytes) > 0

    def test_cover_letter_template_renders_and_contains_expected_text(self):
        ctx = export_rendering.build_cover_letter_docx_context(
            candidate_name="Jane Doe", contact_line="jane@example.com",
            sent_date="August 12, 2026", employer_name="HealthTech Co",
            body_text="Dear Hiring Manager,\n\nI am writing to apply.",
        )
        file_bytes = export_rendering.render_docx(export_templates.COVER_LETTER_TEMPLATE_FILE, ctx)
        from io import BytesIO
        doc = Document(BytesIO(file_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "HealthTech Co" in full_text
        assert "I am writing to apply." in full_text


class TestBuildApplicationPackZip:
    def test_zip_contains_two_independent_docx_entries(self):
        import zipfile
        from io import BytesIO

        zip_bytes = export_rendering.build_application_pack_zip(
            cv_docx=b"fake-cv-docx-bytes", cover_letter_docx=b"fake-cover-letter-bytes",
        )
        zf = zipfile.ZipFile(BytesIO(zip_bytes))
        assert set(zf.namelist()) == {"tailored_cv.docx", "cover_letter.docx"}
        assert zf.read("tailored_cv.docx") == b"fake-cv-docx-bytes"
        assert zf.read("cover_letter.docx") == b"fake-cover-letter-bytes"
