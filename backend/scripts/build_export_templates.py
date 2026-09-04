"""Builds the .docx export templates used by app/services/export_rendering.py.

One-time build script, same precedent as scripts/ingest_esco.py/
ingest_onet.py — the output (app/templates/exports/*.docx) is committed
to the repo, this script isn't run at request time. Real .docx files
with docxtpl (Jinja2-over-docx) tags written as plain paragraph text —
docxtpl's own documented pattern is one control tag ({% %}) per
paragraph, which is what this script produces.

Single-column, no tables/text-boxes/multi-column sections, real
selectable text — deliberately consistent with what
app/extraction/ats_check.py already checks for, so a CV exported through
these templates trivially passes this product's own ATS-readiness check.

Two CV layouts share the exact same tag structure (so they consume the
identical Jinja context from export_rendering.build_cv_docx_context) and
differ only in spacing/font-size:
  - "standard": generous spacing, larger headings — for shorter CVs.
  - "compact": tighter spacing/smaller fonts — for longer CVs that would
    otherwise spill past a reasonable page count.

The actual visual/typographic polish here is a first pass, not a final
design — flagged in the Sprint 5 plan as a real design task, not
something a build script alone can finish.

Usage:
    python scripts/build_export_templates.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

_OUT_DIR = Path(__file__).resolve().parent.parent / "app" / "templates" / "exports"

_ACCENT = RGBColor(0x1F, 0x3A, 0x5F)


def _set_style(doc: Document, *, base_pt: int, heading_pt: int, name_pt: int, spacing_after: int) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(base_pt)
    normal.paragraph_format.space_after = Pt(spacing_after)


def _tag_paragraph(doc: Document, text: str) -> None:
    """A paragraph whose entire content is one Jinja control tag
    ({% ... %}) — docxtpl's documented pattern for block tags that need
    to wrap other paragraphs without leaving stray empty paragraphs in
    the rendered output."""
    doc.add_paragraph(text)


def _build_cv_template(*, base_pt: int, heading_pt: int, name_pt: int, spacing_after: int) -> Document:
    doc = Document()
    _set_style(doc, base_pt=base_pt, heading_pt=heading_pt, name_pt=name_pt, spacing_after=spacing_after)

    # ── Header: name + contact line, both optional (never fabricated —
    # this CV parser doesn't extract a candidate name/phone/location
    # today, only an account email when the export is account-owned; see
    # export_rendering.py for the honest, graceful-omission handling) ──
    _tag_paragraph(doc, "{% if candidate_name %}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{ candidate_name }}")
    run.bold = True
    run.font.size = Pt(name_pt)
    run.font.color.rgb = _ACCENT
    _tag_paragraph(doc, "{% endif %}")

    _tag_paragraph(doc, "{% if contact_line %}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{ contact_line }}")
    run.font.size = Pt(base_pt)
    _tag_paragraph(doc, "{% endif %}")

    # ── Body: one heading per block, then that block's content ──
    _tag_paragraph(doc, "{% for block in blocks %}")

    p = doc.add_paragraph()
    run = p.add_run("{{ block.heading }}")
    run.bold = True
    run.font.size = Pt(heading_pt)
    run.font.color.rgb = _ACCENT
    p.paragraph_format.space_before = Pt(spacing_after)

    # Single-paragraph content (summary, skills)
    _tag_paragraph(doc, "{% if block.paragraph %}")
    doc.add_paragraph("{{ block.paragraph }}")
    _tag_paragraph(doc, "{% endif %}")

    # Flat line list (education)
    _tag_paragraph(doc, "{% if block.lines %}")
    _tag_paragraph(doc, "{% for line in block.lines %}")
    doc.add_paragraph("{{ line }}", style="List Bullet")
    _tag_paragraph(doc, "{% endfor %}")
    _tag_paragraph(doc, "{% endif %}")

    # Headed sub-items with their own bullet lists (experience, projects).
    # Named "entries", not "items" — a dict context key named "items"
    # collides with Python dict's own .items() method under Jinja2's
    # attribute-then-subscript lookup order (confirmed live: it silently
    # resolves to the bound method, not the data, and blows up trying to
    # iterate it) — same class of footgun as "keys"/"values".
    _tag_paragraph(doc, "{% if block.entries %}")
    _tag_paragraph(doc, "{% for item in block.entries %}")
    _tag_paragraph(doc, "{% if item.header %}")
    p = doc.add_paragraph()
    run = p.add_run("{{ item.header }}")
    run.bold = True
    _tag_paragraph(doc, "{% endif %}")
    _tag_paragraph(doc, "{% for bullet in item.bullets %}")
    doc.add_paragraph("{{ bullet }}", style="List Bullet")
    _tag_paragraph(doc, "{% endfor %}")
    _tag_paragraph(doc, "{% endfor %}")
    _tag_paragraph(doc, "{% endif %}")

    _tag_paragraph(doc, "{% endfor %}")

    return doc


def _build_cover_letter_template() -> Document:
    doc = Document()
    _set_style(doc, base_pt=11, heading_pt=11, name_pt=11, spacing_after=8)

    _tag_paragraph(doc, "{% if candidate_name %}")
    doc.add_paragraph("{{ candidate_name }}")
    _tag_paragraph(doc, "{% endif %}")
    _tag_paragraph(doc, "{% if contact_line %}")
    doc.add_paragraph("{{ contact_line }}")
    _tag_paragraph(doc, "{% endif %}")

    doc.add_paragraph("{{ sent_date }}")
    doc.add_paragraph()

    _tag_paragraph(doc, "{% if employer_name %}")
    doc.add_paragraph("{{ employer_name }}")
    _tag_paragraph(doc, "{% endif %}")
    doc.add_paragraph()

    _tag_paragraph(doc, "{% for paragraph in body_paragraphs %}")
    doc.add_paragraph("{{ paragraph }}")
    _tag_paragraph(doc, "{% endfor %}")

    return doc


def main() -> None:
    _OUT_DIR.mkdir(parents=True, exist_ok=True)

    standard = _build_cv_template(base_pt=11, heading_pt=14, name_pt=20, spacing_after=8)
    standard.save(_OUT_DIR / "cv_standard.docx")

    compact = _build_cv_template(base_pt=10, heading_pt=12, name_pt=16, spacing_after=4)
    compact.save(_OUT_DIR / "cv_compact.docx")

    cover_letter = _build_cover_letter_template()
    cover_letter.save(_OUT_DIR / "cover_letter.docx")

    print(f"Wrote 3 export templates to {_OUT_DIR}")


if __name__ == "__main__":
    main()
